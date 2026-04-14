"""
extract_text.py
---------------
Extracts plain text from PDF, DOCX, and TXT resume files.
Each format has its own private helper so the logic stays clean.
"""

import os
import logging

import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


def extract_text(file_path: str) -> str:
    """
    Extract plain text from a resume file.

    Supported formats: .pdf  .docx  .txt

    Args:
        file_path: Absolute path to the resume file.

    Returns:
        Extracted text as a single UTF-8 string.

    Raises:
        FileNotFoundError: File does not exist on disk.
        ValueError:        File format is not supported.
        RuntimeError:      Low-level parsing error inside the chosen library.
    """
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()

    dispatch = {
        ".pdf":  _extract_pdf,
        ".docx": _extract_docx,
        ".txt":  _extract_txt,
    }

    handler = dispatch.get(ext)
    if handler is None:
        supported = ", ".join(dispatch.keys())
        raise ValueError(
            f"Unsupported file format '{ext}'. Supported: {supported}"
        )

    return handler(file_path)


# ── Private helpers ────────────────────────────────────────────────────────────

def _extract_pdf(file_path: str) -> str:
    """
    Use pdfplumber to pull text from every page of a PDF.
    Pages with no selectable text (e.g. scanned images) are skipped with a warning.
    """
    text_parts: list[str] = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text.strip())
                else:
                    logger.warning("Page %d has no extractable text (possibly a scanned image)", i)
    except Exception as exc:
        raise RuntimeError(f"Failed to parse PDF '{file_path}': {exc}") from exc

    return "\n".join(text_parts)


def _extract_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file paragraph by paragraph.
    Tables are also read so skill keywords inside tables are captured.
    """
    try:
        doc = Document(file_path)

        # Collect paragraph text
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        # Collect text from tables (skills are often in table cells)
        table_cells: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        table_cells.append(cell_text)

        return "\n".join(paragraphs + table_cells)
    except Exception as exc:
        raise RuntimeError(f"Failed to parse DOCX '{file_path}': {exc}") from exc


def _extract_txt(file_path: str) -> str:
    """Read a plain-text file, trying UTF-8 then falling back to latin-1."""
    for encoding in ("utf-8", "latin-1"):
        try:
            with open(file_path, "r", encoding=encoding) as fh:
                return fh.read()
        except UnicodeDecodeError:
            continue
    raise RuntimeError(f"Could not decode '{file_path}' with UTF-8 or latin-1.")
